import docx
from pdfminer.high_level import extract_text
import spacy
import re
from sklearn.feature_extraction.text import CountVectorizer
from sklearn.metrics.pairwise import cosine_similarity
def extract_text_from_docx(docx_path):
    doc = docx.Document(docx_path)
    full_text = []
    for paragraph in doc.paragraphs:
        full_text.append(paragraph.text)
    return '\n'.join(full_text)
def extract_text_from_docx(docx_path):
    doc = docx.Document(docx_path)
    full_text = []
    for paragraph in doc.paragraphs:
        full_text.append(paragraph.text)
    return '\n'.join(full_text)
def extract_text(file_path):
    if file_path.endswith('.pdf'):
        return extract_text_from_pdf(file_path)
    elif file_path.endswith('.docx'):
        return extract_text_from_docx(file_path)
    else:
        raise ValueError("Unsupported file format")
# Function to extract information from resume text
def extract_info(text):
    doc = nlp(text)
    name = None
    contact_info = []
    skills = []

    # Extract Name
    for ent in doc.ents:
        if ent.label_ == 'PERSON':
            name = ent.text
            break

    # Extract Contact Info
    contact_info = re.findall(r'\b[\w\.-]+@[\w\.-]+\.\w{2,4}\b', text)
    contact_info += re.findall(r'\b\d{10}\b', text)

    # Extract Skills
    skillset = ['Python', 'Java', 'Machine Learning', 'SQL', 'JavaScript']  # Add more as needed
    skills = [skill for skill in skillset if skill.lower() in text.lower()]

    return {
        'name': name,
        'contact_info': contact_info,
        'skills': skills
    }

# Function to match resume to job description
def match_resume_to_job(resume_text, job_description):
    documents = [resume_text, job_description]
    count_vectorizer = CountVectorizer()
    count_matrix = count_vectorizer.fit_transform(documents)

    match_percentage = cosine_similarity(count_matrix)[0][1] * 100
    return match_percentage

# Function to analyze a resume
def analyze_resume(file_path, job_description):
    text = extract_text(file_path)
    info = extract_info(text)
    match = match_resume_to_job(text, job_description)
    
    return {
        'info': info,
        'match': match
    }
import docx
import fitz  # PyMuPDF
import spacy
import re
from sklearn.feature_extraction.text import CountVectorizer
from sklearn.metrics.pairwise import cosine_similarity

# Load Spacy model
nlp = spacy.load("en_core_web_sm")

# Function to extract text from DOCX files
def extract_text_from_docx(docx_path):
    doc = docx.Document(docx_path)
    full_text = []
    for paragraph in doc.paragraphs:
        full_text.append(paragraph.text)
    return '\n'.join(full_text)

# Function to extract text from PDF files using PyMuPDF
def extract_text_from_pdf(pdf_path):
    try:
        doc = fitz.open(pdf_path)
        text = ""
        for page in doc:
            text += page.get_text()
        return text
    except Exception as e:
        print(f"Error extracting text from PDF: {e}")
        return ""

# Function to determine which extractor to use
def extract_text(file_path):
    if file_path.endswith('.pdf'):
        return extract_text_from_pdf(file_path)
    elif file_path.endswith('.docx'):
        return extract_text_from_docx(file_path)
    else:
        raise ValueError("Unsupported file format")

# Function to extract information from resume text
def extract_info(text):
    doc = nlp(text)
    name = None
    contact_info = []
    skills = []

    # Extract Name
    for ent in doc.ents:
        if ent.label_ == 'PERSON':
            name = ent.text
            break

    # Extract Contact Info
    contact_info = re.findall(r'\b[\w\.-]+@[\w\.-]+\.\w{2,4}\b', text)
    contact_info += re.findall(r'\b\d{10}\b', text)

    # Extract Skills
    skillset = ['Python', 'Java', 'Machine Learning', 'SQL', 'JavaScript']  # Add more as needed
    skills = [skill for skill in skillset if skill.lower() in text.lower()]

    return {
        'name': name,
        'contact_info': contact_info,
        'skills': skills
    }

# Function to match resume to job description
def match_resume_to_job(resume_text, job_description):
    documents = [resume_text, job_description]
    count_vectorizer = CountVectorizer()
    count_matrix = count_vectorizer.fit_transform(documents)

    match_percentage = cosine_similarity(count_matrix)[0][1] * 100
    return match_percentage

# Function to analyze a resume
def analyze_resume(file_path, job_description):
    text = extract_text(file_path)
    info = extract_info(text)
    match = match_resume_to_job(text, job_description)
    
    return {
        'info': info,
        'match': match
    }

# Main function to test the analysis
if __name__ == "__main__":
    job_desc = """
    We are looking for a Data Scientist with experience in Python, Machine Learning, and SQL,html,css.javascript.
    """
    resume_path = "/content/Keerthika Resume.pdf"  # or "/content/Keerthy_resume.docx"
    result = analyze_resume(resume_path, job_desc)
    
    print(f"Name: {result['info']['name']}")
    print(f"Contact: {result['info']['contact_info']}")
    print(f"Skills: {result['info']['skills']}")
    print(f"Match Percentage: {result['match']}%")
